from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth import get_user_model
from django import forms 
from django.http import HttpResponse
from django.contrib.auth.forms import UserCreationForm, UserChangeForm
from django.core.paginator import Paginator
from docx import Document
# from reportlab.lib import colors
import PyPDF2
from .forms import ProfilePictureForm, ReferenceDocumentForm, DatasetDocumentForm, BlogPostForm, BlogCommentForm, BlogCategoryForm
from io import BytesIO
# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter
# from reportlab.lib.units import inch
import re
from difflib import SequenceMatcher
from datetime import date
from django.utils import timezone
from django.conf import settings
from django.db.models import Count, Q
import unicodedata
import os

from .models import ReferenceDocument, CustomUser, PlagiarismHistory, DatasetDocument, TrainedDatasetModel, BlogPost, BlogComment, BlogCategory
from plagiarismchecker.algorithm import fileSimilarity
from plagiarismchecker.algorithm.main import train_tfidf_model, infer_similarity, containment_similarity
from datasets import load_dataset
import threading # For thread-safe caching if needed

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.colors import black, darkred 

class CustomUserCreationForm(UserCreationForm):
    """
    A custom form for creating new users. It extends Django's UserCreationForm
    and explicitly lists fields from the CustomUser model.
    """
    # Override the Meta class entirely to avoid issues with swapped user model
    class Meta:
        model = CustomUser
        # Explicitly list all the fields needed in the registration form.
        # UserCreationForm handles the password fields ('password', 'password2')
        # Ensure 'username', 'email', 'first_name', 'last_name' exist in CustomUser.
        fields = ('username', 'first_name', 'last_name', 'email')

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # Apply the 'interactive-input' class to all form fields
        # This will ensure styling from signup.html is applied to all fields
        for field_name in self.fields:
            self.fields[field_name].widget.attrs.update({'class': 'form-control interactive-input'})
            
        # Make first_name, last_name, and email required, if they aren't by default
        # To ensure they are filled during signup.
        self.fields['first_name'].required = True
        self.fields['last_name'].required = True
        self.fields['email'].required = True

# This view renders webpage.html and is intended for logged-in users.
@login_required
def webpage_content_view(request):
    """
    Renders webpage.html, acting as the primary page for logged-in users.
    """
    return render(request, 'pc/webpage.html')

def public_blog_list(request):
    """Public read-only blog list for unauthenticated access."""
    posts = BlogPost.objects.filter(status='published').order_by('-published_at')

    # public search/category filters
    query = request.GET.get('q')
    if query:
        posts = posts.filter(
            Q(title__icontains=query) |
            Q(content__icontains=query) |
            Q(tags__icontains=query) |
            Q(author__username__icontains=query)
        )

    category_slug = request.GET.get('category')
    if category_slug:
        posts = posts.filter(category__slug=category_slug)

    # Add approved comment count to each post
    for post in posts:
        post.approved_comments_count = post.comments.filter(is_approved=True).count()

    paginator = Paginator(posts, 6)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Re-add approved comment count after pagination
    for post in page_obj:
        post.approved_comments_count = post.comments.filter(is_approved=True).count()

    categories = BlogCategory.objects.all()

    context = {
        'posts': page_obj,
        'categories': categories,
        'current_category': category_slug,
        'search_query': query,
        'is_public_view': True,
    }
    return render(request, 'pc/blog_list.html', context)

def public_blog_detail(request, slug):
    """Public read-only blog detail for unauthenticated access."""
    post = get_object_or_404(BlogPost, slug=slug, status='published')

    # Increment view count for public views as well
    post.views += 1
    post.save(update_fields=['views'])

    comments = post.comments.filter(is_approved=True)

    # Auto-compute similar posts for public users too (same as authenticated users)
    auto_similar_results = []
    try:
        # Compare current post against all other published posts with content
        candidate_posts = BlogPost.objects.filter(status='published').exclude(pk=post.pk).exclude(content='')
        temp = []
        for other in candidate_posts:
            try:
                # Use containment similarity for fast computation
                score = containment_similarity(post.content or '', other.content or '')
            except Exception:
                score = 0.0
            if score > 5.0:  # Only show results with >5% similarity
                temp.append({
                    'post': other,
                    'similarity_percentage': round(float(score), 2),
                })
        # Sort by highest similarity and keep top 5 to keep UI concise
        temp.sort(key=lambda r: r['similarity_percentage'], reverse=True)
        auto_similar_results = temp[:5]
    except Exception:
        auto_similar_results = []

    context = {
        'post': post,
        'comments': comments,
        'comment_form': None,
        'is_public_view': True,
        'plagiarism_results': None,
        'total_posts_checked': 0,
        'auto_similar_results': auto_similar_results,
    }
    return render(request, 'pc/blog_detail.html', context)
@login_required
def profile_view(request):
    """
    Displays the user's profile and their plagiarism check history,
    and handles profile picture updates.
    """
    user = request.user
    history = PlagiarismHistory.objects.filter(user=user).order_by('-checked_at')

    if request.method == 'POST':
        # Instantiate the form with POST data and files, and the user instance
        pic_form = ProfilePictureForm(request.POST, request.FILES, instance=user)
        if pic_form.is_valid():
            pic_form.save()
            messages.success(request, 'Your profile picture was successfully updated!')
            return redirect('profile') # Redirect to refresh the page and show new pic
        else:
            # Display form errors using Django messages
            for field, errors in pic_form.errors.items():
                for error in errors:
                    messages.error(request, f"Profile Picture Error: {error}")
    else:
        # For GET requests, instantiate an empty form with the user instance for initial data
        pic_form = ProfilePictureForm(instance=user)

    # Retrieve messages to pass to the template
    current_messages = messages.get_messages(request)

    context = {
        'user': user,
        'history': history,
        'pic_form': pic_form, # Pass the form instance to the template
        'messages': current_messages, # Pass messages for display
    }
    return render(request, 'registration/profile_view.html', context)

class CustomUserChangeForm(UserChangeForm):
    """
    A custom form for updating existing users. Useful for admin or profile editing.
    """
    class Meta:
        model = CustomUser
        fields = ('username', 'first_name', 'last_name', 'email', 'is_active', 'is_staff', 'is_superuser')
        # Add other fields that might need to allow changing
_cached_plagiarism_dataset = None
_dataset_lock = threading.Lock()
def get_plagiarism_dataset():
    global _cached_plagiarism_dataset
    if _cached_plagiarism_dataset is None:
        with _dataset_lock:
            if _cached_plagiarism_dataset is None: # Double-check locking
                try:
                    print("Loading MIT-PLAGAIRISM-DETECTION-DATASET for the first time...")
                    _cached_plagiarism_dataset = load_dataset("jatinmehra/MIT-PLAGAIRISM-DETECTION-DATASET", split="train")
                    print("Dataset loaded.")
                except Exception as dataset_error:
                    print(f"Failed to load HF dataset: {dataset_error}")
                    _cached_plagiarism_dataset = None
    return _cached_plagiarism_dataset

# --- Utility: robust normalization for exact-substring matching ---
def _normalize_for_exact_match(value: str) -> str:
    try:
        text = value.lower()
        # Collapse all whitespace (spaces, tabs, newlines) to single spaces
        text = re.sub(r"\s+", " ", text)
        # Remove zero-width or special spaces
        text = text.replace("\u200b", "").strip()
        return text
    except Exception:
        return value.lower().strip()

# Call get_plagiarism_dataset() to retrieve it.
# Helper Functions (assuming these are defined elsewhere or provided by user's project)
def extract_text_from_file(uploaded_file):
    """
    Extracts text content from uploaded .docx, .pdf, or .txt files and performs robust cleaning.
    """
    text_content = ""
    file_name = uploaded_file.name

    try:
        file_extension = file_name.split('.')[-1].lower()
        if file_extension == 'txt':
            text_content = uploaded_file.read().decode('utf-8')
        elif file_extension == 'pdf':
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                text_content += page.extract_text() or ""
        elif file_extension == 'docx':
            document = Document(uploaded_file)
            # Extract paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text:
                    text_content += paragraph.text + "\n"
            # Extract all table cell text as well
            try:
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                text_content += cell_text + "\n"
            except Exception as e:
                # Non-fatal; some documents may not have tables or may raise if malformed
                print(f"DOCX table extraction skipped: {e}")
        else:
            raise ValueError(f"Unsupported file type: {file_name}. Please upload a .txt, .pdf, or .docx file.")
    except Exception as e:
        print(f"Error reading file {file_name}: {e}")
        raise ValueError(f"Could not read file: {file_name}. Error: {e}")

    # --- ROBUST TEXT CLEANING (APPLIED HERE FOR FILE INPUT) ---
    if text_content:
        # Normalize Unicode characters (e.g., ligatures, accented characters)
        text_content = unicodedata.normalize('NFKC', text_content)
        # Replace common non-breaking space with regular space
        text_content = text_content.replace('\xa0', ' ')
        # Replace any sequence of whitespace characters (including newlines, tabs, multiple spaces) with a single space
        text_content = re.sub(r'\s+', ' ', text_content)
        # Remove leading/trailing whitespace
        text_content = text_content.strip()
    # --- END ROBUST TEXT CLEANING ---

    return text_content

def get_reference_text():
    """
    Fetches combined text from all ReferenceDocument objects in the database.
    """
    all_reference_texts = ReferenceDocument.objects.all()
    combined_text = ""
    for doc in all_reference_texts:
        text = (doc.content or "").strip()
        if not text and getattr(doc, 'document_file', None):
            try:
                with open(doc.document_file.path, 'rb') as f:
                    text = extract_text_from_file(f)
            except Exception as e:
                print(f"Failed to read document_file for ReferenceDocument id={doc.id}: {e}")
                text = ""
        combined_text += (text or "") + "\n"
    return combined_text.strip()

def find_plagiarized_spans(text1, text2, min_match_length=5):
    """
    Identify matching spans between two texts using word-level matching.
    Returns a list of (start_char_index, end_char_index) tuples in text1.
    min_match_length is specified in words (not characters).
    """
    # Tokenize to words and keep character span for each token
    word_regex = re.compile(r'\b\w+\b')
    tokens1 = []
    token1_spans = []
    for m in word_regex.finditer(text1.lower()):
        tokens1.append(m.group(0))
        token1_spans.append((m.start(), m.end()))

    tokens2 = []
    for m in word_regex.finditer(text2.lower()):
        tokens2.append(m.group(0))

    # If either side has no tokens, nothing to do
    if not tokens1 or not tokens2:
        return []

    matcher = SequenceMatcher(None, tokens1, tokens2)
    spans = []
    for block in matcher.get_matching_blocks():
        # block: (i, j, n) where tokens1[i:i+n] == tokens2[j:j+n]
        if block.size >= max(1, int(min_match_length)):
            start_token_index = block.a
            end_token_index = block.a + block.size - 1
            # Map back to character span using token1_spans
            start_char = token1_spans[start_token_index][0]
            end_char = token1_spans[end_token_index][1]
            spans.append((start_char, end_char))
    return spans

def count_plagiarized_words(text_content, spans):
    """
    Counts the total number of words within the given spans in text_content,
    merging overlapping or adjacent spans to prevent overcounting.
    """
    if not spans:
        return 0

    # Ensure all spans are in dictionary format and sort them by their start index
    # This defensive conversion helps prevent errors if the span format is inconsistent
    processed_spans = []
    for s in spans:
        try:
            processed_spans.append({
                'start': int(s['start']),
                'end': int(s['end'])
            })
        except (TypeError, ValueError) as e:
            print(f"Error converting span to int: {s}. Skipping this span. Error: {e}")
            continue # Skip malformed spans

    if not processed_spans:
        return 0

    processed_spans.sort(key=lambda x: x['start'])

    # Merge overlapping and adjacent spans
    merged_spans = []
    current_start = processed_spans[0]['start']
    current_end = processed_spans[0]['end']

    for i in range(1, len(processed_spans)):
        next_start = processed_spans[i]['start']
        next_end = processed_spans[i]['end']

        # If the next span overlaps or is immediately adjacent (e.g., end of 1st is 10, start of 2nd is 11)
        if next_start <= current_end + 1:
            current_end = max(current_end, next_end)
        else:
            # Current merged span is complete, add it to list
            merged_spans.append({'start': current_start, 'end': current_end})
            # Start a new merged span with the next span's values
            current_start = next_start
            current_end = next_end
    
    # Add the last merged span after the loop finishes
    merged_spans.append({'start': current_start, 'end': current_end})

    plag_word_count = 0
    for span_dict in merged_spans:
        start = span_dict['start']
        end = span_dict['end']
        
        # Defensive check against non-integer indices right before slicing
        if not isinstance(start, int) or not isinstance(end, int):
            print(f"ERROR: Non-integer indices found after merging: start={start} ({type(start)}), end={end} ({type(end)})")
            continue # Skip this problematic span
            
        segment = text_content[start:end]
        # Count words in the segment using your preferred word definition
        plag_word_count += len(re.findall(r'\b\w+\b', segment))
        
    return plag_word_count

def convert_spans_to_dicts(spans_list):
    """
    Converts a list of (start, end) tuples to a list of {'start': start, 'end': end} dictionaries.
    If spans_list already contains dictionaries, it returns it as is.
    This handles cases where old session data might contain tuple-based spans, preventing
    "tuple indices must be integers or slices, not str" errors.
    """
    if not spans_list:
        return []
    # Check if the first element is a tuple (indicating old format)
    if isinstance(spans_list[0], tuple):
        return [{'start': s[0], 'end': s[1]} for s in spans_list]
    return spans_list # Already in dictionary format


def generate_plagiarism_report_pdf(text, spans, percent, web_results=None):
    try:
        print(f"DEBUG: Starting PDF generation - text length: {len(text)}, spans: {len(spans) if spans else 0}")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom style for highlighted text
        highlight_style = ParagraphStyle(
            name='HighlightStyle',
            parent=styles['Normal'],
            textColor=darkred, # Using imported darkred
            fontName='Helvetica',
            fontSize=10,
            leading=12, # Line spacing
        )

        # Custom style for normal text
        normal_style = ParagraphStyle(
            name='NormalStyle',
            parent=styles['Normal'],
            textColor=black, # Using imported black
            fontName='Helvetica',
            fontSize=10,
            leading=12,
        )
        
        # Add a style for titles/headers
        title_style = ParagraphStyle(
            name='TitleStyle',
            parent=styles['h1'],
            alignment=TA_LEFT,
            fontName='Helvetica-Bold',
            fontSize=16,
            spaceAfter=14
        )

        header_style = ParagraphStyle(
            name='HeaderStyle',
            parent=styles['h2'],
            alignment=TA_LEFT,
            fontName='Helvetica-Bold',
            fontSize=12,
            spaceAfter=6
        )

        story = [] # This list will hold all the ReportLab flowables (paragraphs, spaces, etc.)

        # Title
        story.append(Paragraph("Plagiarism Report", title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Overall Plagiarism Percentage
        story.append(Paragraph(f"Overall Plagiarism: {percent:.2f}%", header_style))
        story.append(Spacer(1, 0.2 * inch))
        
        story.append(Paragraph("Original Text (Plagiarized sections highlighted):", header_style))
        story.append(Spacer(1, 0.1 * inch))

        # Ensure spans are in dictionary format and sort them by their start index
        converted_spans = []
        for s in spans:
            try:
                # Defensive conversion to int in case values are floats or strings
                converted_spans.append({'start': int(s['start']), 'end': int(s['end'])})
            except (TypeError, ValueError) as e:
                print(f"Skipping malformed span during PDF generation: {s}. Error: {e}")
                continue
        converted_spans.sort(key=lambda x: x['start'])

        # Build RML (ReportLab Markup Language) string for highlighting
        # RML allows embedding HTML-like tags for styling
        rml_text_parts = []
        last_idx = 0
        for span in converted_spans:
            # Non-plagiarized segment
            if span['start'] > last_idx:
                # Escape special characters for RML if the text contains '&', '<', '>'
                rml_text_parts.append(text[last_idx:span['start']].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
            
            # Plagiarized segment with highlight
            plag_segment = text[span['start']:span['end']]
            plag_segment = plag_segment.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            rml_text_parts.append(f"<font color='darkred'>{plag_segment}</font>")
            last_idx = span['end']
        
        # Remaining non-plagiarized text
        if last_idx < len(text):
            rml_text_parts.append(text[last_idx:].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))

        # Join all segments and create a Paragraph
        # Replace newlines with <br/> for ReportLab Paragraphs to force line breaks
        final_text_content_rml = "".join(rml_text_parts).replace('\n', '<br/>')
        story.append(Paragraph(final_text_content_rml, normal_style))
        story.append(Spacer(1, 0.3 * inch))

        # Web Results (if available)
        if web_results and web_results.get('output'):
            story.append(Paragraph("Web Plagiarism Matches:", header_style))
            story.append(Spacer(1, 0.1 * inch))
            
            for link, data in web_results['output'].items():
                title = data.get('title', 'N/A')
                snippet = data.get('snippet', 'N/A')
                
                story.append(Paragraph(f"<b>Title:</b> {title}", normal_style))
                # Links should ideally be clickable, but basic text is fine for now
                story.append(Paragraph(f"<b>Link:</b> <font color='blue'>{link.replace('&', '&amp;')}</font>", normal_style))
                story.append(Paragraph(f"<b>Snippet:</b> {snippet[:300].replace('&', '&amp;')}. . .", normal_style)) # Truncate and escape snippet
                story.append(Spacer(1, 0.15 * inch))

        # Build the PDF document
        print("DEBUG: Building PDF document...")
        doc.build(story)
        buffer.seek(0)
        pdf_data = buffer.getvalue()
        print(f"DEBUG: PDF generation completed, size: {len(pdf_data)} bytes")
        return pdf_data
        
    except Exception as e:
        print(f"DEBUG: PDF generation error: {e}")
        import traceback
        traceback.print_exc()
        raise Exception(f"PDF generation failed: {str(e)}")

def count_plagiarized_words(text_content, spans):
    """
    Counts the total number of words within the given spans in text_content,
    merging overlapping or adjacent spans to prevent overcounting.
    """
    if not spans:
        return 0

    # Ensure all spans are in dictionary format and sort them by their start index
    processed_spans = []
    for s in spans:
        try:
            processed_spans.append({
                'start': int(s['start']),
                'end': int(s['end'])
            })
        except (TypeError, ValueError) as e:
            print(f"Error converting span to int: {s}. Skipping this span. Error: {e}")
            continue

    if not processed_spans:
        return 0

    processed_spans.sort(key=lambda x: x['start'])

    # Merge overlapping and adjacent spans
    merged_spans = []
    current_start = processed_spans[0]['start']
    current_end = processed_spans[0]['end']

    for i in range(1, len(processed_spans)):
        next_start = processed_spans[i]['start']
        next_end = processed_spans[i]['end']

        # If the next span overlaps or is immediately adjacent (e.g., end of 1st is 10, start of 2nd is 11)
        if next_start <= current_end + 1:
            current_end = max(current_end, next_end)
        else:
            # Current merged span is complete, add it to list
            merged_spans.append({'start': current_start, 'end': current_end})
            # Start a new merged span with the next span's values
            current_start = next_start
            current_end = next_end
    
    # Add the last merged span after the loop finishes
    merged_spans.append({'start': current_start, 'end': current_end})

    plag_word_count = 0
    for span_dict in merged_spans:
        start = span_dict['start']
        end = span_dict['end']
        
        # Defensive check against non-integer indices right before slicing
        if not isinstance(start, int) or not isinstance(end, int):
            print(f"ERROR: Non-integer indices found after merging: start={start} ({type(start)}), end={end} ({type(end)})")
            continue
            
        segment = text_content[start:end]
        # Count words in the segment using your preferred word definition
        plag_word_count += len(re.findall(r'\b\w+\b', segment))
        
    return plag_word_count

def generate_comparison_report_pdf(report_data):
    """Generates a PDF report for two-document comparison."""
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    p.setFont("Helvetica-Bold", 16)
    p.drawString(inch, height - inch, "Document Comparison Report")

    p.setFont("Helvetica", 12)
    p.drawString(inch, height - inch - 30, f"Similarity Percentage: {report_data['similarity_percent']}%")
    p.drawString(inch, height - inch - 50, f"Matching Words: {report_data['matching_words']}")
    p.drawString(inch, height - inch - 70, f"Total Unique Words: {report_data['total_words']}")

    y_position = height - inch - 100
    
    p.drawString(inch, y_position, "Document 1 Content:")
    y_position -= 15
    for line in report_data['text1'].split('\n'):
        if y_position < inch:
            p.showPage()
            p.setFont("Helvetica", 12)
            y_position = height - inch
        p.drawString(inch, y_position, line[:100]) # Truncate long lines for PDF display
        y_position -= 14

    y_position -= 30 # Space between docs

    p.drawString(inch, y_position, "Document 2 Content:")
    y_position -= 15
    for line in report_data['text2'].split('\n'):
        if y_position < inch:
            p.showPage()
            p.setFont("Helvetica", 12)
            y_position = height - inch
        p.drawString(inch, y_position, line[:100]) # Truncate long lines for PDF display
        y_position -= 14

    p.save()
    buffer.seek(0)
    return buffer

# View functions

def homepage(request):
    """
    Redirects logged-in users to the community blog; otherwise, displays frontpage.html.
    This function will be mapped to the root URL ('/').
    """
    if request.user.is_authenticated:
        # If logged in, redirect to the community blog
        return redirect('blog_list')
    else:
        # If not logged in, render frontpage.html directly
        return render(request, 'pc/frontpage.html')

def signup(request):
    if request.method == 'POST':
        # Pass request.FILES if CustomUserCreationForm handles profile pictures
        form = CustomUserCreationForm(request.POST, request.FILES)
        if form.is_valid():
            user = form.save()
            # Log the user in immediately after successful signup
            login(request, user)
            messages.success(request, f'Account created successfully! Welcome, {user.username}!')
            # Redirect to the community blog after signup AND login
            return redirect('blog_list')
        else:
            # Display detailed form errors as messages
            for field, errors in form.errors.items():
                for error in errors:
                    # 'non_field_errors' means errors not tied to a specific input field
                    if field == '__all__':
                        messages.error(request, f"Error: {error}")
                    else:
                        messages.error(request, f"{field.capitalize()}: {error}")
    else:
        form = CustomUserCreationForm()
    # Ensure correct template path
    return render(request, 'registration/signup.html', {'form': form})

# --- Login View ---
def user_login(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        
        # Authenticate the user using the provided credentials
        # (Assuming CustomUser uses 'username' as USERNAME_FIELD)
        user = authenticate(request, username=username, password=password)

        if user is not None:
            # If authentication is successful, log the user in
            login(request, user)
            messages.success(request, f'Welcome back, {user.username}!')
            # Redirect to the community blog after successful login
            return redirect('blog_list')
        else:
            # If authentication fails, display an error message
            messages.error(request, 'Invalid username or password. Please try again.')
    # Ensure correct template path
    return render(request, 'registration/login_page.html')

@login_required
def user_logout(request):
    logout(request)
    messages.info(request, 'You have been logged out.')
    return redirect('home')

# Duplicate profile_view removed (the earlier one with picture update stays)

# Plagiarism Detection View (Single Document)
def index(request):
    """
    Handles plagiarism detection (text, file, or web input) and PDF report generation,
    now including checks against a Hugging Face dataset.
    Results are stored in the session for display and PDF download across requests.
    """
    # Initialize variables to ensure they are always defined for context
    overall_plagiarism_percent = None # Will store the max percent from all checks
    total_words = None
    plag_words = None
    error = None
    q_text_input = ''
    all_plagiarized_spans = [] # Will collect spans from all sources
    file_name = None    # To store the uploaded file's name
    web_results = None  # To store web search results
    hf_dataset_results = None # To store Hugging Face dataset comparison results
    plagiarism_sources = [] # To store detailed source information

    if request.method == 'POST':
        action = request.POST.get('action')
        download_pdf_requested = request.POST.get('download_pdf') == '1'

        input_text_content = ""

        # PDF-related variables; populated from session later if PDF is requested
        text_for_pdf = ''
        percent_for_pdf = None
        spans_for_pdf = []
        web_results_for_pdf = None
        # hf_dataset_results_for_pdf = None # Will decide how to integrate this into PDF later

        try:
            # Debug logging for action
            print(f"DEBUG: Action received: {action}")
            print(f"DEBUG: POST data keys: {list(request.POST.keys())}")
            print(f"DEBUG: FILES data keys: {list(request.FILES.keys())}")
            print(f"DEBUG: Download PDF requested: {download_pdf_requested}")
            
            # Handle session clearing request first
            if action == 'clear_session':
                print("DEBUG: Clearing session data...")
                # Clear all plagiarism check related session data
                session_keys_to_clear = [
                    'index_last_check_text',
                    'index_last_check_percent', 
                    'index_last_check_plag_words',
                    'index_last_check_total_words',
                    'index_last_check_spans',
                    'q_text_input',
                    'web_search_results',
                    'hf_dataset_results',
                    'active_trained_dataset',
                    'plagiarism_sources'
                ]
                
                for key in session_keys_to_clear:
                    if key in request.session:
                        del request.session[key]
                        print(f"DEBUG: Cleared session key: {key}")
                
                print("DEBUG: Session cleared successfully")
                messages.success(request, "Results cleared successfully!")
                return redirect('detect')
            
            # Handle PDF download request (uses existing session data)
            elif action == 'download_pdf' or download_pdf_requested:
                print("DEBUG: PDF download requested")
                
                # Retrieve all necessary data from session for PDF generation
                text_for_pdf = request.session.get('index_last_check_text', '')
                percent_for_pdf = request.session.get('index_last_check_percent')
                spans_for_pdf = request.session.get('index_last_check_spans', [])
                web_results_for_pdf = request.session.get('web_search_results', None)
                
                print(f"DEBUG: PDF data - Text length: {len(text_for_pdf)}, Percent: {percent_for_pdf}, Spans: {len(spans_for_pdf) if spans_for_pdf else 0}")
                
                # Check if we have the minimum required data
                if not text_for_pdf:
                    raise ValueError("No text data found for PDF generation. Please run a plagiarism check first.")
                
                if percent_for_pdf is None:
                    # If no percentage is available, set it to 0
                    percent_for_pdf = 0.0
                    print("DEBUG: No percentage found, setting to 0.0")

                spans_for_pdf = convert_spans_to_dicts(spans_for_pdf) if spans_for_pdf else []
                
                try:
                    print("DEBUG: Generating PDF...")
                    pdf = generate_plagiarism_report_pdf(
                        text_for_pdf,
                        spans_for_pdf,
                        percent_for_pdf,
                        web_results_for_pdf
                    )
                    
                    if pdf:
                        print(f"DEBUG: PDF generated successfully, size: {len(pdf)} bytes")
                        response = HttpResponse(pdf, content_type='application/pdf')
                        response['Content-Disposition'] = 'attachment; filename="plagiarism_report.pdf"'
                        response['Content-Length'] = len(pdf)
                        return response
                    else:
                        raise ValueError("PDF generation returned empty result")
                        
                except Exception as pdf_error:
                    print(f"DEBUG: PDF generation failed: {pdf_error}")
                    raise ValueError(f"Failed to generate PDF report: {str(pdf_error)}")
            
            elif action == 'check_text' or action == 'check_web' or action == 'check_file':
                if action == 'check_text' or action == 'check_web':
                    input_text_content = request.POST.get('q', '').strip()
                    if not input_text_content:
                        raise ValueError("Please enter some text to check for plagiarism!")
                    q_text_input = input_text_content
                    print(f"DEBUG: Text input length: {len(input_text_content)}")
                elif action == 'check_file':
                    print("DEBUG: Starting file upload processing...")
                    uploaded_file = request.FILES.get('docfile')
                    if not uploaded_file:
                        raise ValueError("Please select a file to check for plagiarism!")
                    
                    file_name = uploaded_file.name
                    file_size = uploaded_file.size
                    print(f"DEBUG: Processing file: {file_name}, Size: {file_size} bytes")
                    
                    # Validate file type
                    allowed_extensions = ['txt', 'pdf', 'docx']
                    file_extension = file_name.split('.')[-1].lower() if '.' in file_name else ''
                    print(f"DEBUG: File extension detected: {file_extension}")
                    if file_extension not in allowed_extensions:
                        raise ValueError(f"Unsupported file type: .{file_extension}. Please upload a .txt, .pdf, or .docx file.")
                    
                    print("DEBUG: Starting text extraction from file...")
                    try:
                        input_text_content = extract_text_from_file(uploaded_file)
                        print(f"DEBUG: Text extraction successful. Length: {len(input_text_content) if input_text_content else 0}")
                        
                        if not input_text_content or len(input_text_content.strip()) == 0:
                            raise ValueError(f"Could not extract text from {file_name}. The file might be empty or corrupted.")
                        
                        # Show first 100 characters for debugging
                        preview = input_text_content[:100].replace('\n', ' ').replace('\r', ' ')
                        print(f"DEBUG: Text preview: {preview}...")
                        
                    except Exception as extraction_error:
                        print(f"DEBUG: Text extraction failed: {extraction_error}")
                        raise ValueError(f"Failed to extract text from {file_name}: {str(extraction_error)}")
                    
                    print(f"DEBUG: File processing complete. Final text length: {len(input_text_content)}")
            else:
                # If no valid action (text/file/web check) and not a PDF download, something is wrong
                if not download_pdf_requested:
                    raise ValueError(f"Invalid action received: {action}. Please try again.")

            # --- Initialize overall plagiarism metrics ---
            overall_plagiarism_percent = 0.0
            all_plagiarized_spans = [] # Re-initialize for fresh run
            
            # --- Perform plagiarism checks based on the action ---
            if input_text_content: # Ensure we have text to check
                total_words = len(re.findall(r'\b\w+\b', input_text_content))
                print(f"DEBUG: Starting plagiarism checks for action '{action}' with {total_words} words")

                # --- 1. Local/File Reference Check with Source Detection ---
                if action == 'check_text' or action == 'check_file':
                    print(f"DEBUG: Starting local reference check for action: {action}")
                    try:
                        reference_text = get_reference_text()
                        if reference_text:
                            print(f"DEBUG: Reference text length: {len(reference_text)}")
                            print("DEBUG: Calculating file similarity...")
                            percent_local = round(fileSimilarity.findFileSimilarity(input_text_content, reference_text), 2)
                            print(f"DEBUG: File similarity calculation complete: {percent_local}%")
                            
                            print("DEBUG: Finding plagiarized spans...")
                            spans_local = find_plagiarized_spans(input_text_content, reference_text)
                            print(f"DEBUG: Found {len(spans_local) if spans_local else 0} plagiarized spans")
                            
                            all_plagiarized_spans.extend(convert_spans_to_dicts(spans_local))
                            
                            # Extract detailed source information from individual reference documents
                            if percent_local > 5:  # Only process if significant similarity
                                ref_docs = ReferenceDocument.objects.all()
                                for doc in ref_docs:
                                    doc_similarity = fileSimilarity.findFileSimilarity(input_text_content, doc.content)
                                    if doc_similarity > 5:
                                        doc_spans = find_plagiarized_spans(input_text_content, doc.content, min_match_length=3)
                                        if doc_spans:
                                            # Extract matching text segments
                                            matching_segments = []
                                            for start, end in doc_spans[:3]:  # Limit to top 3 matches per source
                                                user_text = input_text_content[start:end].strip()
                                                # Find corresponding text in source document
                                                source_spans = find_plagiarized_spans(doc.content, input_text_content, min_match_length=3)
                                                source_text = ""
                                                if source_spans:
                                                    src_start, src_end = source_spans[0]
                                                    source_text = doc.content[src_start:src_end].strip()
                                                
                                                if user_text and len(user_text) > 10:  # Only meaningful segments
                                                    matching_segments.append({
                                                        'user_text': user_text,
                                                        'source_text': source_text or user_text,
                                                        'char_start': start,
                                                        'char_end': end
                                                    })
                                            
                                            if matching_segments:
                                                plagiarism_sources.append({
                                                    'source_title': doc.title,
                                                    'source_author': doc.author,
                                                    'similarity_percentage': round(doc_similarity, 2),
                                                    'matching_segments': matching_segments,
                                                    'source_type': 'Reference Document'
                                                })
                            
                            overall_plagiarism_percent = max(overall_plagiarism_percent, percent_local)
                            print(f"DEBUG: Local similarity complete: {percent_local}%, Overall: {overall_plagiarism_percent}%")
                            print(f"DEBUG: Found {len(plagiarism_sources)} plagiarism sources from reference documents")
                            # plag_words will be calculated after all spans are collected
                        else:
                            print("DEBUG: No reference text found in database")
                    except Exception as e:
                        print(f"DEBUG: Local reference check failed: {e}")
                        import traceback
                        traceback.print_exc()
                        # Continue with other checks even if local check fails

                # --- 2. Custom Trained Dataset Check (TF-IDF) across ALL trained datasets ---
                try:
                    selected_dataset = request.GET.get('dataset')
                    trained_qs = TrainedDatasetModel.objects.all()
                    if selected_dataset:
                        trained_qs = trained_qs.filter(dataset_name=selected_dataset)

                    best_percent = 0.0
                    best_meta = None
                    for trained in trained_qs:
                        model_dir = os.path.dirname(trained.vectorizer_path)
                        percent_custom_tfidf = infer_similarity(input_text_content, model_dir)
                        corpus_docs = DatasetDocument.objects.filter(dataset_name=trained.dataset_name).values_list('content', flat=True)
                        corpus_text = "\n".join([c or '' for c in corpus_docs])
                        # Direct exact-substring check after strong normalization
                        q_norm = _normalize_for_exact_match(input_text_content)
                        corpus_norm = _normalize_for_exact_match(corpus_text)
                        if q_norm and q_norm in corpus_norm:
                            percent_custom = 100.0
                            # Apply exact span over the entire query so word count reflects match
                            all_plagiarized_spans.append({'start': 0, 'end': len(input_text_content)})
                        else:
                            percent_containment = containment_similarity(input_text_content, corpus_text, n=5)
                            percent_custom = max(percent_custom_tfidf, percent_containment)
                        percent_custom = round(percent_custom, 2)
                        if percent_custom > best_percent:
                            best_percent = percent_custom
                            best_meta = {
                                'dataset_name': trained.dataset_name,
                                'trained_at': trained.trained_at.isoformat() if trained.trained_at else None,
                                'percent_custom': percent_custom,
                            }

                    if best_meta is not None:
                        overall_plagiarism_percent = max(overall_plagiarism_percent, best_meta['percent_custom'])
                        request.session['active_trained_dataset'] = best_meta
                except Exception as e:
                    print(f"Custom model inference failed: {e}")

                # --- 3. Hugging Face Dataset Check---
                try:
                    dataset = get_plagiarism_dataset()
                    if dataset is not None:
                        print("DEBUG: Hugging Face dataset loaded successfully")
                        current_hf_similarities = {} # Maps dataset document ID/title to similarity percent
                        current_hf_max_percent = 0.0

                        # Set the limit for the number of documents to process from the dataset
                        limit_docs = 20 # adjust this number (e.g., 50, 200, 500)

                        for doc_index, document in enumerate(dataset):
                            #Check if the limit has been reached, and break the loop if it has
                            if doc_index >= limit_docs:
                                print(f"DEBUG: Stopped processing Hugging Face dataset after {limit_docs} documents.")
                                break 
                                
                            # Assuming the dataset has a 'text' or 'source_text' key.
                            dataset_doc_text = document.get('text', document.get('source_text')) # Try common keys
                            dataset_doc_id = document.get('id', f"DatasetDoc_{doc_index}") # Use 'id' or generate one

                            if dataset_doc_text:
                                sim_percent_hf = round(fileSimilarity.findFileSimilarity(input_text_content, dataset_doc_text), 2)
                                if sim_percent_hf > 0.0: # Only store if there's some similarity
                                    current_hf_similarities[dataset_doc_id] = sim_percent_hf
                                    current_hf_max_percent = max(current_hf_max_percent, sim_percent_hf)
                                    
                                    # Get spans from this dataset document match
                                    hf_spans = find_plagiarized_spans(input_text_content, dataset_doc_text)
                                    all_plagiarized_spans.extend(convert_spans_to_dicts(hf_spans))
                                    
                        hf_dataset_results = {
                            'similarities': current_hf_similarities,
                            'max_percent_found': current_hf_max_percent
                        }
                        overall_plagiarism_percent = max(overall_plagiarism_percent, current_hf_max_percent)
                        print(f"DEBUG: HF dataset max similarity: {current_hf_max_percent}%")
                        # plag_words will be calculated after all spans are collected
                    else:
                        print("DEBUG: Hugging Face dataset not available")
                except Exception as e:
                    print(f"DEBUG: Hugging Face dataset check failed: {e}")
                    # Continue with other checks even if HF dataset check fails

                # --- 4. Web Search Check (enhanced with containment & page fetch) ---
                if action == 'check_web':
                    try:
                        print("DEBUG: Starting web search")
                        print(f"DEBUG: Input text for web search (first 100 chars): {input_text_content[:100]}...")
                        from plagiarismchecker.algorithm.webSearch import searchWeb
                        import requests
                        output = {}
                        c = {}
                        print("DEBUG: Calling searchWeb function...")
                        output, c, web_error_code = searchWeb(input_text_content, output, c)
                        print(f"DEBUG: Web search completed. Error code: {web_error_code}")
                        print(f"DEBUG: Web search output keys: {list(output.keys()) if output else 'No output'}")
                        print(f"DEBUG: Web search similarities: {c}")
                        
                        if web_error_code == 0:
                            best_web_percent = 0.0
                            enhanced_output = {}
                            any_exact_web = False
                            web_spans_collected = []
                            
                            for link, data in output.items():
                                title = data.get('title', '') or ''
                                snippet = data.get('snippet', '') or ''
                                # Base cosine score (0..1) -> percent
                                cosine_score = float(c.get(link, 0)) * 100.0
                                # Snippet containment
                                snip_score = containment_similarity(input_text_content, snippet, n=5)
                                
                                # Try page fetch for better containment and span detection
                                page_score = 0.0
                                page_text = ""
                                try:
                                    resp = requests.get(link, timeout=6)
                                    if resp.ok:
                                        html = resp.text
                                        html = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.IGNORECASE)
                                        html = re.sub(r"<style[\s\S]*?</style>", " ", html, flags=re.IGNORECASE)
                                        page_text = re.sub(r"<[^>]+>", " ", html)
                                        page_text = re.sub(r'\s+', ' ', page_text).strip()  # Clean up whitespace
                                        
                                        if _normalize_for_exact_match(input_text_content) in _normalize_for_exact_match(page_text):
                                            page_score = 100.0
                                            any_exact_web = True
                                            # For exact matches, mark entire text as plagiarized
                                            web_spans_collected.append({'start': 0, 'end': len(input_text_content)})
                                        else:
                                            page_score = containment_similarity(input_text_content, page_text, n=5)
                                            # Find specific matching spans for partial matches
                                            if page_score > 5:  # Only extract spans for significant matches
                                                page_spans = find_plagiarized_spans(input_text_content, page_text, min_match_length=3)
                                                web_spans_collected.extend(convert_spans_to_dicts(page_spans))
                                except Exception as e:
                                    print(f"DEBUG: Failed to fetch page {link}: {e}")
                                    # Fall back to snippet-based span detection
                                    if snip_score > 5:
                                        snippet_spans = find_plagiarized_spans(input_text_content, snippet, min_match_length=3)
                                        web_spans_collected.extend(convert_spans_to_dicts(snippet_spans))
                                
                                link_best = max(cosine_score, snip_score, page_score)
                                best_web_percent = max(best_web_percent, link_best)
                                enhanced_output[link] = {
                                    'title': title,
                                    'snippet': snippet,
                                    'score': round(link_best, 2),
                                    'page_text_available': bool(page_text)
                                }
                            
                            # Add all web spans to the overall collection
                            all_plagiarized_spans.extend(web_spans_collected)
                            print(f"DEBUG: Added {len(web_spans_collected)} web spans to overall collection")
                            
                            # Extract detailed source information from web results
                            for link, data in enhanced_output.items():
                                if data['score'] > 5:  # Only process significant matches
                                    # Try to get page content for better source text extraction
                                    source_text = ""
                                    try:
                                        resp = requests.get(link, timeout=6)
                                        if resp.ok:
                                            html = resp.text
                                            html = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.IGNORECASE)
                                            html = re.sub(r"<style[\s\S]*?</style>", " ", html, flags=re.IGNORECASE)
                                            source_text = re.sub(r"<[^>]+>", " ", html)
                                            source_text = re.sub(r'\s+', ' ', source_text).strip()
                                    except Exception:
                                        source_text = data['snippet']  # Fallback to snippet
                                    
                                    if source_text:
                                        # Find matching segments between user text and web source
                                        web_source_spans = find_plagiarized_spans(input_text_content, source_text, min_match_length=3)
                                        matching_segments = []
                                        
                                        for start, end in web_source_spans[:3]:  # Limit to top 3 matches per source
                                            user_text = input_text_content[start:end].strip()
                                            # Find corresponding text in web source
                                            source_spans = find_plagiarized_spans(source_text, input_text_content, min_match_length=3)
                                            web_source_text = ""
                                            if source_spans:
                                                src_start, src_end = source_spans[0]
                                                web_source_text = source_text[src_start:src_end].strip()
                                            
                                            if user_text and len(user_text) > 10:  # Only meaningful segments
                                                matching_segments.append({
                                                    'user_text': user_text,
                                                    'source_text': web_source_text or user_text,
                                                    'char_start': start,
                                                    'char_end': end
                                                })
                                        
                                        if matching_segments:
                                            plagiarism_sources.append({
                                                'source_title': data['title'] or 'Web Source',
                                                'source_author': 'Web Content',
                                                'source_url': link,
                                                'similarity_percentage': data['score'],
                                                'matching_segments': matching_segments,
                                                'source_type': 'Web Source'
                                            })
                            
                            web_results = {'output': enhanced_output, 'similarities': c}
                            overall_plagiarism_percent = max(overall_plagiarism_percent, round(best_web_percent, 2))
                            print(f"DEBUG: Web check complete. Best web percent: {best_web_percent}%, Overall: {overall_plagiarism_percent}%")
                            print(f"DEBUG: Found {len([s for s in plagiarism_sources if s['source_type'] == 'Web Source'])} web plagiarism sources")
                    except Exception as e:
                        print(f"Web search temporarily unavailable: {e}")

                # This line is good as it ensures consistency, but the above changes
                # make `all_plagiarized_spans` consistently dictionary-based from its origin.
                # all_plagiarized_spans = convert_spans_to_dicts(all_plagiarized_spans) # Can keep for robustness

                # Calculate plag_words based on ALL collected and MERGED spans
                # IMPORTANT: Ensure count_plagiarized_words function
                # properly handles these dictionary-based spans and merges them to avoid overcounting.
                plag_words = count_plagiarized_words(input_text_content, all_plagiarized_spans)
                print(f"DEBUG: Total words: {total_words}, Plagiarized words: {plag_words}")
                print(f"DEBUG: Number of spans found: {len(all_plagiarized_spans)}")
                
                # Fix inconsistency: Calculate final similarity based on plagiarized words vs total words
                if plag_words == 0:
                    print("DEBUG: No plagiarized words found, setting similarity percentage to 0")
                    word_based_percent = 0.0
                elif total_words and total_words > 0:
                    # Percentage of words identified as plagiarized
                    word_based_percent = round((plag_words / total_words) * 100, 2)
                    print(f"DEBUG: Word-based plagiarism percentage: {word_based_percent}%")
                else:
                    word_based_percent = 0.0
                
                # For web checks, use the higher of algorithm-based or word-based percentage
                if action == 'check_web':
                    # Web checks should use word-based calculation as primary metric
                    overall_plagiarism_percent = word_based_percent
                    print(f"DEBUG: Web check - using word-based percentage: {word_based_percent}%")
                else:
                    # For text/file checks, use word-based percentage as the most accurate measure
                    overall_plagiarism_percent = word_based_percent
                
                print(f"DEBUG: Final overall plagiarism percentage: {overall_plagiarism_percent}%")
            # --- Save all results to PlagiarismHistory for database logging ---    
            if input_text_content:
                history_query_text = ""
                content_preview_length = 200
                if file_name:
                    history_query_text = f"File: {file_name} - Content: {input_text_content[:content_preview_length]}"
                else:
                    history_query_text = f"Text Input: {input_text_content[:content_preview_length]}"
                
                if len(history_query_text) > 500:
                    history_query_text = history_query_text[:497] + "..."

                if request.user.is_authenticated:
                    PlagiarismHistory.objects.create(
                        user=request.user,
                        query_text=history_query_text,
                        result_percentage=overall_plagiarism_percent # Use the final overall percent
                    )

            # --- Store all final results in session for immediate display and future PDF download ---
            request.session['index_last_check_text'] = input_text_content
            request.session['index_last_check_percent'] = overall_plagiarism_percent
            request.session['index_last_check_plag_words'] = plag_words
            request.session['index_last_check_total_words'] = total_words
            request.session['index_last_check_spans'] = all_plagiarized_spans # All collected spans
            request.session['q_text_input'] = input_text_content
            request.session['web_search_results'] = web_results # Full web results data
            request.session['hf_dataset_results'] = hf_dataset_results # New: Full HF dataset results data
            request.session['plagiarism_sources'] = plagiarism_sources # Store detailed source information

            # Clear session data not needed for current display or future PDF
            # This logic needs to be careful to not clear results from other sources if action changes.
            # It's better to store all results and manage which ones are displayed or used for PDF.
            # The current approach of explicitly clearing is prone to issues if multiple checks are done.
            # Instead, just store what's relevant to the current check type and let the context decide.
            # The code now stores ALL relevant results and leaves it to the template/PDF to decide what to show.
            # The previous "clearing" logic for web vs local was removed for simplicity and correctness.




        except ValueError as ve:
            error = str(ve)
            messages.error(request, error)
        except Exception as e:
            error = f"Uh oh! An unexpected error popped up: {e}. I'm sorry about that!"
            messages.error(request, error)
        
        return redirect('detect')

    # --- GET request handling (or after redirect from POST) ---
    # Retrieve results from session for display on the page
    # Using get() instead of pop() to preserve data for PDF downloads
    overall_plagiarism_percent = request.session.get('index_last_check_percent', None)
    total_words = request.session.get('index_last_check_total_words', None)
    plag_words = request.session.get('index_last_check_plag_words', None)
    all_plagiarized_spans = request.session.get('index_last_check_spans', [])
    web_results = request.session.get('web_search_results', None)
    hf_dataset_results = request.session.get('hf_dataset_results', None) # New: Retrieve HF results
    custom_model_info = request.session.get('active_trained_dataset', None)

    q_text_input = request.session.get('q_text_input', '')
    # Keep text data in session for PDF generation
    # request.session.pop('index_last_check_text', None)  # Commented out to preserve for PDF

    all_plagiarized_spans = convert_spans_to_dicts(all_plagiarized_spans)

    error_messages = messages.get_messages(request)
    current_error = None
    for msg in error_messages:
        if msg.tags == 'error':
            current_error = str(msg)
            break
            
    # Retrieve plagiarism sources from session if available
    session_plagiarism_sources = request.session.get('plagiarism_sources', [])
            
    context = {
        'percent': overall_plagiarism_percent, # Renamed for clarity in context
        'total_words': total_words,
        'plag_words': plag_words,
        'error': current_error,
        'q_text_input': q_text_input,
        'spans': all_plagiarized_spans,
        'web_results': web_results,
        'hf_dataset_results': hf_dataset_results, # New: Pass HF results to template
        'custom_model_info': custom_model_info,
        'plagiarism_sources': session_plagiarism_sources, # Add plagiarism sources to context
    }
    return render(request, 'pc/index.html', context)

# Document Comparison View (Two Documents)
@login_required
def twofilecompare1(request): # This is the function that handles doc_compare.html
    """
    Handles comparison of two documents (text or file input) and PDF report generation.
    """
    text_result = None
    file_result = None
    error = None
    matching_words = None
    total_words = None

    q1_content = ''
    q2_content = ''

    if request.method == 'POST':
        download_pdf_requested = request.POST.get("download_pdf") == "1"
        form_type = request.POST.get('form_type')

        # Initialize PDF-related variables for two-file comparison to ensure they are always defined
        text1_for_pdf = ''
        text2_for_pdf = ''
        percent_for_pdf = None
        total_words_for_pdf = None
        matching_words_for_pdf = None

        try:
            similarity_percent = 0.0
            matching_words_val = 0
            total_words_val = 0
            
            if form_type == 'text_compare': # For text comparison
                q1_content = request.POST.get('q1', '').strip()
                q2_content = request.POST.get('q2', '').strip()

                if not q1_content or not q2_content:
                    raise ValueError("Hey! You need to put text in BOTH boxes for me to compare them.")
                
                # from plagiarismchecker.algorithm import fileSimilarity # Already imported at top
                similarity_percent = round(fileSimilarity.findFileSimilarity(q1_content, q2_content), 2)
                
                input_words_set = set(re.findall(r'\b\w+\b', q1_content.lower()))
                compare_words_set = set(re.findall(r'\b\w+\b', q2_content.lower()))
                matching_words_val = len(input_words_set.intersection(compare_words_set))
                total_words_val = len(input_words_set.union(compare_words_set))

                # Store results for text comparison in specific session variables
                request.session['compare_text_last_content1'] = q1_content
                request.session['compare_text_last_content2'] = q2_content
                request.session['compare_text_last_percent'] = similarity_percent
                request.session['compare_text_last_total_words'] = total_words_val
                request.session['compare_text_last_matching_words'] = matching_words_val
                # Clear file compare session data if text compare was done
                request.session.pop('compare_file_last_percent', None)
                request.session.pop('compare_file_last_content1', None)
                request.session.pop('compare_file_last_content2', None)
                request.session.pop('compare_file_last_total_words', None)
                request.session.pop('compare_file_last_matching_words', None)


            elif form_type == 'file_compare': # For file comparison
                file1 = request.FILES.get("docfile1")
                file2 = request.FILES.get("docfile2")

                if not file1 or not file2:
                    raise ValueError("Whoops! I need BOTH files uploaded for comparison.")
                
                q1_content = extract_text_from_file(file1)
                q2_content = extract_text_from_file(file2)

                if not q1_content or not q2_content:
                    raise ValueError("Hmm, I couldn't get any text from one or both files. Are they empty or a weird format?")
                
                # from plagiarismchecker.algorithm import fileSimilarity # Already imported at top
                similarity_percent = round(fileSimilarity.findFileSimilarity(q1_content, q2_content), 2)

                input_words_set = set(re.findall(r'\b\w+\b', q1_content.lower()))
                compare_words_set = set(re.findall(r'\b\w+\b', q2_content.lower()))
                matching_words_val = len(input_words_set.intersection(compare_words_set))
                total_words_val = len(input_words_set.union(compare_words_set))

                # Store results for file comparison in distinct session variables
                request.session['compare_file_last_content1'] = q1_content
                request.session['compare_file_last_content2'] = q2_content
                request.session['compare_file_last_percent'] = similarity_percent
                request.session['compare_file_last_total_words'] = total_words_val
                request.session['compare_file_last_matching_words'] = matching_words_val
                # Clear text compare session data if file compare was done
                request.session.pop('compare_text_last_percent', None)
                request.session.pop('compare_text_last_content1', None)
                request.session.pop('compare_text_last_content2', None)
                request.session.pop('compare_text_last_total_words', None)
                request.session.pop('compare_text_last_matching_words', None)

            else:
                if not download_pdf_requested: # Only raise error if it's not a PDF download request with missing form_type
                    raise ValueError("No valid comparison input provided (text or files).")

            # Save to PlagiarismHistory if a comparison was made (not just PDF download) ---
            if request.user.is_authenticated and form_type in ['text_compare', 'file_compare'] and (q1_content or q2_content):
                # Ensure query_text doesn't exceed max_length (e.g., 500 characters)
                # Combine a snippet of both documents for the history entry
                combined_query_text = f"Doc 1: {q1_content[:200]}... vs Doc 2: {q2_content[:200]}"
                if len(combined_query_text) > 500: # Model's query_text max_length
                    combined_query_text = combined_query_text[:497] + '...'

                PlagiarismHistory.objects.create(
                    user=request.user,
                    query_text=combined_query_text,
                    result_percentage=similarity_percent
                )
            # --- END NEW: PlagiarismHistory saving ---

            if download_pdf_requested:
                # Assign values from session to the initialized variables
                if 'compare_text_last_percent' in request.session:
                    text1_for_pdf = request.session.get('compare_text_last_content1', '')
                    text2_for_pdf = request.session.get('compare_text_last_content2', '')
                    percent_for_pdf = request.session.get('compare_text_last_percent')
                    total_words_for_pdf = request.session.get('compare_text_last_total_words')
                    matching_words_for_pdf = request.session.get('compare_text_last_matching_words')
                elif 'compare_file_last_percent' in request.session:
                    text1_for_pdf = request.session.get('compare_file_last_content1', '')
                    text2_for_pdf = request.session.get('compare_file_last_content2', '')
                    percent_for_pdf = request.session.get('compare_file_last_percent')
                    total_words_for_pdf = request.session.get('compare_file_last_total_words')
                    matching_words_for_pdf = request.session.get('compare_file_last_matching_words')
                else:
                    raise ValueError("Please run a document comparison first before trying to download the PDF report.")

                if text1_for_pdf and text2_for_pdf and percent_for_pdf is not None:
                    report_data = {
                        'text1': text1_for_pdf,
                        'text2': text2_for_pdf,
                        'similarity_percent': percent_for_pdf,
                        'total_words': total_words_for_pdf,
                        'matching_words': matching_words_for_pdf,
                    }
                    pdf = generate_comparison_report_pdf(report_data)
                    response = HttpResponse(pdf, content_type='application/pdf')
                    response['Content-Disposition'] = 'attachment; filename="document_comparison_report.pdf"'
                    return response
                else:
                    raise ValueError("Failed to retrieve complete comparison data for PDF report. Please run a comparison again.")

        except ValueError as ve:
            error = str(ve)
        except Exception as e:
            error = f"Uh oh! An unexpected error popped up: {e}. I'm sorry about that!"
        
        if error:
            messages.error(request, error)
        
        return redirect('compare') # Redirect to the GET request of the same view

    # GET request handling or after redirect
    # First, check for text comparison results in session
    if 'compare_text_last_percent' in request.session:
        text_result = request.session.pop('compare_text_last_percent')
        matching_words = request.session.pop('compare_text_last_matching_words', None)
        total_words = request.session.pop('compare_text_last_total_words', None)
        q1_content = request.session.pop('compare_text_last_content1', '')
        q2_content = request.session.pop('compare_text_last_content2', '')
        file_result = None # Ensure file_result is None if text results are present
    # Then check for file comparison results
    elif 'compare_file_last_percent' in request.session:
        file_result = request.session.pop('compare_file_last_percent')
        matching_words = request.session.pop('compare_file_last_matching_words', None)
        total_words = request.session.pop('compare_file_last_total_words', None)
        q1_content = request.session.pop('compare_file_last_content1', '')
        q2_content = request.session.pop('compare_file_last_content2', '')
        text_result = None # Ensure text_result is None if file results are present
    else:
        # No results in session, keep everything as None or default
        text_result = None
        file_result = None
        matching_words = None
        total_words = None
        q1_content = ''
        q2_content = ''

    error_messages = messages.get_messages(request)
    current_error = None
    for msg in error_messages:
        if msg.tags == 'error':
            current_error = str(msg)
            break

    context = {
        'text_result': text_result,
        'file_result': file_result,
        'matching_words': matching_words,
        'total_words': total_words,
        'q1_content': q1_content,
        'q2_content': q2_content,
        'error': current_error,
    }
    
    return render(request, 'pc/doc_compare.html', context)

def check_plagiarism(request):
    result = None
    checked_text = None
    if request.method == 'POST':
        checked_text = request.POST.get('text')
        if not checked_text:
            messages.error(request, 'Please enter text to check.')
            context = {
                'result': result, 
                'checked_text': checked_text,
                'plagiarism_sources': []
            }
            return render(request, 'pc/index.html', context)
        if not request.user.is_authenticated:
            session_data = request.session.get('anonymous_plagiarism_data', {})
            today = date.today().isoformat()

            if session_data.get('last_check_date') != today:
                session_data['word_count_today'] = 0
                session_data['last_check_date'] = today

            word_count = len(checked_text.split())

            if session_data['word_count_today'] + word_count > 1000:
                messages.error(request, 'Daily word limit (1000 words) exceeded for unauthenticated users. Please log in for unlimited access.')
                context = {
                    'result': result,
                    'checked_text': checked_text,
                    'word_limit_exceeded': True,
                    'current_word_count_anon': session_data['word_count_today'],
                    'daily_limit_anon': 1000,
                }
                # IMPORTANT: Return here if limit exceeded for anon users
                return render(request, 'pc/index.html', context)
            else:
                session_data['word_count_today'] += word_count
                request.session['anonymous_plagiarism_data'] = session_data
                messages.info(request, f"Anonymous check: Words used today: {session_data['word_count_today']}/1000")

        # Enhanced plagiarism checking with source detection
        reference_text = get_reference_text()
        plagiarism_sources = []
        total_similarity = 0.0
        
        if reference_text:
            # Find plagiarized spans and calculate similarity
            from plagiarismchecker.algorithm.fileSimilarity import findFileSimilarity
            similarity_percentage = findFileSimilarity(checked_text, reference_text)
            
            if similarity_percentage > 5:  # Only process if significant similarity
                spans = find_plagiarized_spans(checked_text, reference_text, min_match_length=3)
                
                # Get individual reference documents for source attribution
                ref_docs = ReferenceDocument.objects.all()
                for doc in ref_docs:
                    doc_similarity = findFileSimilarity(checked_text, doc.content)
                    if doc_similarity > 5:
                        doc_spans = find_plagiarized_spans(checked_text, doc.content, min_match_length=3)
                        if doc_spans:
                            # Extract matching text segments
                            matching_segments = []
                            for start, end in doc_spans[:3]:  # Limit to top 3 matches per source
                                user_text = checked_text[start:end].strip()
                                # Find corresponding text in source document
                                source_spans = find_plagiarized_spans(doc.content, checked_text, min_match_length=3)
                                source_text = ""
                                if source_spans:
                                    src_start, src_end = source_spans[0]
                                    source_text = doc.content[src_start:src_end].strip()
                                
                                if user_text and len(user_text) > 10:  # Only meaningful segments
                                    matching_segments.append({
                                        'user_text': user_text,
                                        'source_text': source_text or user_text,
                                        'char_start': start,
                                        'char_end': end
                                    })
                            
                            if matching_segments:
                                plagiarism_sources.append({
                                    'source_title': doc.title,
                                    'source_author': doc.author,
                                    'similarity_percentage': round(doc_similarity, 2),
                                    'matching_segments': matching_segments
                                })
                
                total_similarity = similarity_percentage
            
            result = total_similarity
        else:
            result = 0.0
            messages.warning(request, 'No reference documents found. Please add reference documents to improve plagiarism detection.')

        if request.user.is_authenticated:
            history_query_text = f"Manual Check: {checked_text[:200]}"
            if len(history_query_text) > 500:
                history_query_text = history_query_text[:497] + "..."

            PlagiarismHistory.objects.create(
                user=request.user,
                query_text=history_query_text,
                result_percentage=result
            )
            if plagiarism_sources:
                messages.success(request, f'Plagiarism check complete. Result: {result}% - {len(plagiarism_sources)} source(s) found.')
            else:
                messages.success(request, f'Plagiarism check complete. Result: {result}% - No significant plagiarism detected.')
        else:
            # Message for unauthenticated users after successful check within limits
            if plagiarism_sources:
                messages.success(request, f'Plagiarism check complete. Result: {result}% - {len(plagiarism_sources)} source(s) found.')
            else:
                messages.success(request, f'Plagiarism check complete. Result: {result}% - No significant plagiarism detected.')
    
    context = {
        'result': result, 
        'checked_text': checked_text,
        'plagiarism_sources': plagiarism_sources if 'plagiarism_sources' in locals() else []
    }
    return render(request, 'pc/index.html', context)
            
def add_document(request):
    """
    Handles adding a new ReferenceDocument via a web form.
    """
    if request.method == 'POST':
        form = ReferenceDocument(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Document added successfully!')
            return redirect('webpage') # Redirect to main webpage or a confirmation page
        else:
            # If form is not valid, pass errors back to the template
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field.capitalize()}: {error}")
    else:
        form = ReferenceDocument() # An empty form for GET request

    return render(request, 'add_document.html', {'form': form})


@login_required
def manage_dataset(request):
    """Upload or paste documents into a named dataset for training."""
    if request.method == 'POST':
        form = DatasetDocumentForm(request.POST, request.FILES)
        if form.is_valid():
            doc = form.save(commit=False)
            # Prefer extracting directly from the uploaded file object
            uploaded_file = form.cleaned_data.get('source_file')
            if uploaded_file and not doc.content:
                try:
                    doc.content = extract_text_from_file(uploaded_file)
                except Exception as e:
                    messages.error(request, f"Failed to extract text from file: {e}")
            # Save the model (file will be saved by the model field)
            doc.save()
            messages.success(request, 'Dataset document saved!')
            return redirect('manage_dataset')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field.capitalize()}: {error}")
    else:
        form = DatasetDocumentForm()

    recent_docs = DatasetDocument.objects.order_by('-created_at')[:20]

    # Build status per dataset
    doc_counts = DatasetDocument.objects.values('dataset_name').annotate(num_docs=Count('id'))
    trained_by_name = {m.dataset_name: m for m in TrainedDatasetModel.objects.all()}
    dataset_status = []
    for row in doc_counts:
        name = row['dataset_name']
        trained = trained_by_name.get(name)
        dataset_status.append({
            'dataset_name': name,
            'num_docs': row['num_docs'],
            'trained': bool(trained),
            'trained_at': trained.trained_at if trained else None,
        })

    return render(request, 'pc/manage_dataset.html', {
        'form': form,
        'recent_docs': recent_docs,
        'dataset_status': dataset_status,
    })


@login_required
def train_dataset(request):
    """Train a TF-IDF model for a given dataset name and persist artifacts."""
    dataset_name = request.POST.get('dataset_name', 'default') if request.method == 'POST' else request.GET.get('dataset_name', 'default')
    docs_qs = DatasetDocument.objects.filter(dataset_name=dataset_name).order_by('id')
    documents = [d.content or '' for d in docs_qs]
    documents = [d for d in documents if d.strip()]
    if not documents:
        messages.error(request, f'No documents found for dataset "{dataset_name}".')
        return redirect('manage_dataset')

    # Store artifacts under MEDIA_ROOT/models/<dataset_name>
    model_dir = os.path.join(settings.MEDIA_ROOT, 'models', dataset_name)
    vectorizer_path, matrix_path, index_path = train_tfidf_model(documents, model_dir)

    TrainedDatasetModel.objects.update_or_create(
        dataset_name=dataset_name,
        defaults={
            'vectorizer_path': vectorizer_path,
            'matrix_path': matrix_path,
            'doc_index_path': index_path,
        }
    )
    messages.success(request, f'Trained TF-IDF model for dataset "{dataset_name}" with {len(documents)} documents.')
    return redirect('manage_dataset')

def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    result = 0 # Default value
    plagiarism_matches = [] # Initialize empty list for matches

    if request.POST['q1'] != '' and request.POST['q2'] != '':
        print("Got both the texts")
        # Update to receive both the result and matches
        result, plagiarism_matches = fileSimilarity.findFileSimilarity(
            request.POST['q1'], request.POST['q2']
        )
    result = round(result, 2)

    # Store these in the session for later PDF generation
    request.session['last_checked_text'] = request.POST['q1'] # Assuming q1 is the main text
    request.session['last_plagiarism_matches'] = plagiarism_matches

    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result, 'plagiarism_matches': plagiarism_matches})
def fetch_dataset_rows_from_hf_api(offset=0, length=100):
    import requests
    url = f"https://datasets-server.huggingface.co/rows?dataset=jatinmehra%2FMIT-PLAGAIRISM-DETECTION-DATASET&config=default&split=train&offset={offset}&length={length}"
    try:
        response = requests.get(url)
        response.raise_for_status() # Raise an exception for bad status codes
        data = response.json()
        return data.get('rows', [])
    except requests.exceptions.RequestException as e:
        print(f"Error fetching dataset from API: {e}")
        return []

def is_admin(user):
    return user.is_staff

@login_required
def blog_list(request):
    """Display all published blog posts and user's own draft posts"""
    if request.user.is_staff:
        # Staff can see all posts
        posts = BlogPost.objects.all().order_by('-created_at')
    else:
        # Regular users see published posts and their own draft posts
        posts = BlogPost.objects.filter(
            Q(status='published') | Q(author=request.user)
        ).order_by('-created_at')
    
    # Search functionality
    query = request.GET.get('q')
    if query:
        posts = posts.filter(
            Q(title__icontains=query) |
            Q(content__icontains=query) |
            Q(tags__icontains=query) |
            Q(author__username__icontains=query)
        )
    
    # Category filter
    category_slug = request.GET.get('category')
    if category_slug:
        posts = posts.filter(category__slug=category_slug)
    
    # Add approved comment count to each post
    for post in posts:
        post.approved_comments_count = post.comments.filter(is_approved=True).count()
    
    # Pagination
    paginator = Paginator(posts, 6)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Re-add approved comment count after pagination
    for post in page_obj:
        post.approved_comments_count = post.comments.filter(is_approved=True).count()
    
    categories = BlogCategory.objects.all()
    
    context = {
        'posts': page_obj,
        'categories': categories,
        'current_category': category_slug,
        'search_query': query,
    }
    return render(request, 'pc/blog_list.html', context)

@login_required
def blog_detail(request, slug):
    """Display a single blog post with comments"""
    # Allow authors to view their own draft posts, but only published posts for others
    if request.user.is_staff:
        # Staff can view any post
        post = get_object_or_404(BlogPost, slug=slug)
    else:
        # Regular users can view published posts or their own draft posts
        post = get_object_or_404(
            BlogPost.objects.filter(
                slug=slug
            ).filter(
                Q(status='published') | Q(author=request.user)
            )
        )
    
    # Increment view count
    post.views += 1
    post.save(update_fields=['views'])
    
    # Handle form submissions
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'check_plagiarism':
            # Handle plagiarism detection
            algorithm = request.POST.get('algorithm', 'cosine')
            current_post_content = post.content
            
            # Get all other published posts (excluding the current one)
            other_posts = BlogPost.objects.filter(status='published').exclude(pk=post.pk).exclude(content='')
            
            plagiarism_results = []
            for other_post in other_posts:
                other_post_content = other_post.content
                
                if algorithm == 'cosine':
                    from plagiarismchecker.algorithm.fileSimilarity import findFileSimilarity
                    similarity_percentage = findFileSimilarity(current_post_content, other_post_content)
                elif algorithm == 'containment':
                    from plagiarismchecker.algorithm.main import containment_similarity
                    similarity_percentage = containment_similarity(current_post_content, other_post_content)
                else:
                    from plagiarismchecker.algorithm.fileSimilarity import findFileSimilarity
                    similarity_percentage = findFileSimilarity(current_post_content, other_post_content)
                
                # Only include results with significant similarity (>5%)
                if similarity_percentage > 5:
                    # Find matching text segments between posts
                    matching_spans = find_plagiarized_spans(current_post_content, other_post_content, min_match_length=3)
                    matching_segments = []
                    
                    for start, end in matching_spans[:3]:  # Limit to top 3 matches per source
                        user_text = current_post_content[start:end].strip()
                        # Find corresponding text in source post
                        source_spans = find_plagiarized_spans(other_post_content, current_post_content, min_match_length=3)
                        source_text = ""
                        if source_spans:
                            src_start, src_end = source_spans[0]
                            source_text = other_post_content[src_start:src_end].strip()
                        
                        if user_text and len(user_text) > 10:  # Only meaningful segments
                            matching_segments.append({
                                'user_text': user_text,
                                'source_text': source_text or user_text,
                                'char_start': start,
                                'char_end': end
                            })
                    
                    plagiarism_results.append({
                        'post': other_post,
                        'similarity_percentage': round(similarity_percentage, 2),
                        'algorithm_used': algorithm,
                        'matching_segments': matching_segments
                    })
            
            # Sort by similarity percentage (highest first)
            plagiarism_results.sort(key=lambda x: x['similarity_percentage'], reverse=True)
            
            comment_form = BlogCommentForm()
            # Get approved comments
            comments = post.comments.filter(is_approved=True)
            context = {
                'post': post,
                'comments': comments,
                'comment_form': comment_form,
                'plagiarism_results': plagiarism_results,
                'total_posts_checked': len(other_posts),
                'algorithm_used': algorithm
            }
            return render(request, 'pc/blog_detail.html', context)
        
        elif action == 'add_comment':
            # Handle comment submission
            comment_form = BlogCommentForm(request.POST)
            if comment_form.is_valid():
                comment = comment_form.save(commit=False)
                comment.post = post
                comment.author = request.user
                comment.is_approved = True  # Auto-approve comments
                comment.save()
                messages.success(request, 'Comment submitted successfully!')
                return redirect('blog_detail', slug=slug)
            else:
                messages.error(request, 'Please correct the errors in your comment.')
        else:
            messages.error(request, 'Invalid action specified.')
    else:
        comment_form = BlogCommentForm()
    
    # Get approved comments
    comments = post.comments.filter(is_approved=True)
    
    # Auto-compute similar posts for all users (same as admin panel)
    auto_similar_results = []
    try:
        # Compare current post against all other published posts with content
        candidate_posts = BlogPost.objects.filter(status='published').exclude(pk=post.pk).exclude(content='')
        temp = []
        for other in candidate_posts:
            try:
                # Use containment similarity for fast computation
                score = containment_similarity(post.content or '', other.content or '')
            except Exception:
                score = 0.0
            if score > 5.0:  # Only show results with >5% similarity
                temp.append({
                    'post': other,
                    'similarity_percentage': round(float(score), 2),
                })
        # Sort by highest similarity and keep top 5 to keep UI concise
        temp.sort(key=lambda r: r['similarity_percentage'], reverse=True)
        auto_similar_results = temp[:5]
    except Exception:
        auto_similar_results = []
    
    context = {
        'post': post,
        'comments': comments,
        'comment_form': comment_form,
        'is_public_view': False,
        'auto_similar_results': auto_similar_results,
    }
    return render(request, 'pc/blog_detail.html', context)

@login_required
def check_blog_plagiarism(request):
    """AJAX endpoint to check blog content for plagiarism against other blog posts"""
    if request.method == 'POST':
        import json
        from django.http import JsonResponse
        
        try:
            data = json.loads(request.body)
            content = data.get('content', '').strip()
            
            if not content:
                return JsonResponse({'error': 'No content provided'}, status=400)
            
            # Compare against all existing blog posts (any status), excluding empty content
            existing_posts = BlogPost.objects.exclude(content='')
            
            plagiarism_results = []
            max_similarity = 0
            
            for post in existing_posts:
                try:
                    # Use cosine similarity for comparison
                    from plagiarismchecker.algorithm.main import infer_similarity
                    similarity = infer_similarity(content, post.content)
                    similarity_percent = round(similarity * 100, 2)
                    
                    if similarity_percent > 15:  # Only include significant similarities
                        # Find matching text segments
                        from plagiarismchecker.algorithm.main import find_plagiarized_spans
                        spans = find_plagiarized_spans(content, post.content)
                        
                        matching_segments = []
                        for span in spans[:3]:  # Limit to top 3 segments
                            user_text = content[span['start']:span['end']]
                            matching_segments.append({
                                'user_text': user_text[:200] + ('...' if len(user_text) > 200 else ''),
                                'source_text': span.get('source_text', '')[:200] + ('...' if len(span.get('source_text', '')) > 200 else '')
                            })
                        
                        plagiarism_results.append({
                            'post_title': post.title,
                            'post_author': post.author.username,
                            'similarity_percentage': similarity_percent,
                            'matching_segments': matching_segments
                        })
                        
                        max_similarity = max(max_similarity, similarity_percent)
                
                except Exception as e:
                    print(f"Error checking similarity with post {post.id}: {e}")
                    continue
            
            # Sort by similarity percentage (highest first)
            plagiarism_results.sort(key=lambda x: x['similarity_percentage'], reverse=True)
            
            has_plagiarism = max_similarity >= 60
            block_submission = max_similarity >= 60
            
            return JsonResponse({
                'max_similarity': max_similarity,
                'results': plagiarism_results[:5],  # Return top 5 matches
                'has_plagiarism': has_plagiarism,
                'block_submission': block_submission
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Invalid request method'}, status=405)

@login_required
def blog_create(request):
    """Create a new blog post"""
    if request.method == 'POST':
        # The frontend now sends a 'plagiarism_bypass' field if the user clicks 'Submit Anyway'.
        # We don't need to do anything special with it on the backend other than
        # ensure it doesn't interfere with form validation.
        post_data = request.POST.copy()
        if 'plagiarism_bypass' in post_data:
            del post_data['plagiarism_bypass']

        form = BlogPostForm(post_data, request.FILES, user=request.user)
        if form.is_valid():
            post = form.save(commit=False)
            post.author = request.user
            post.save()

            if post.status == 'pending':
                messages.success(request, 'Blog post submitted for review!')
            else:
                messages.success(request, 'Blog post saved as draft!')
            return redirect('blog_my_posts')
        else:
            import logging
            logging.error(f"Blog post form errors: {form.errors.as_json()}")
            messages.error(request, 'Please correct the errors below.')
    else:
        form = BlogPostForm(user=request.user)

    context = {
        'form': form,
        'title': 'Create New Blog Post',
    }
    return render(request, 'pc/blog_form.html', context)

@login_required
def blog_edit(request, slug):
    """Edit an existing blog post"""
    post = get_object_or_404(BlogPost, slug=slug)

    # Check permissions
    if not request.user.is_staff and post.author != request.user:
        messages.error(request, 'You do not have permission to edit this post.')
        return redirect('blog_list')

    if request.method == 'POST':
        form = BlogPostForm(request.POST, request.FILES, instance=post, user=request.user)
        if form.is_valid():
            content = form.cleaned_data.get('content', '')
            max_similarity = 0

            if content and content.strip():
                # Perform plagiarism check against all other posts, excluding the current one
                from plagiarismchecker.algorithm.main import infer_similarity
                existing_posts = BlogPost.objects.exclude(pk=post.pk).exclude(content__exact='')
                for other_post in existing_posts:
                    similarity = infer_similarity(content, other_post.content)
                    if similarity > max_similarity:
                        max_similarity = similarity

            max_similarity_percent = round(max_similarity * 100, 2)

            form.save()
            messages.success(request, 'Blog post updated successfully!')
            return redirect('blog_detail', slug=post.slug)
    else:
        form = BlogPostForm(instance=post, user=request.user)

    context = {
        'form': form,
        'post': post,
        'title': f'Edit "{post.title}"',
    }
    return render(request, 'pc/blog_form.html', context)

import logging
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Helper function for direct text comparison
def compare_texts(text1, text2):
    """Calculates cosine similarity between two text strings."""
    if not text1 or not text2:
        return 0.0
    try:
        vectorizer = TfidfVectorizer().fit_transform([text1, text2])
        vectors = vectorizer.toarray()
        similarity = cosine_similarity(vectors)
        return similarity[0, 1]
    except Exception as e:
        logging.error(f"Error comparing texts: {e}")
        return 0.0

@login_required
def blog_my_posts(request):
    """Display user's own blog posts with plagiarism info."""
    user_posts = BlogPost.objects.filter(author=request.user).order_by('-created_at')
    all_posts = BlogPost.objects.all()

    posts_with_plagiarism_info = []
    for post in user_posts:
        max_similarity = 0
        most_similar_post = None

        # Compare against all posts except itself
        comparison_posts = all_posts.exclude(pk=post.pk)

        if post.content and comparison_posts.exists():
            for other_post in comparison_posts:
                if other_post.content:
                    # Use the new direct comparison function
                    similarity = compare_texts(post.content, other_post.content)
                    if similarity > max_similarity:
                        max_similarity = similarity
                        most_similar_post = other_post
        
        similarity_percentage = round(max_similarity * 100, 2)
        logging.info(f"Plagiarism check for post '{post.title}': {similarity_percentage}%")

        posts_with_plagiarism_info.append({
            'post': post,
            'max_similarity': similarity_percentage,
            'most_similar_post': most_similar_post
        })

    context = {
        'posts_with_info': posts_with_plagiarism_info,
    }
    return render(request, 'pc/blog_my_posts.html', context)

@user_passes_test(is_admin)
def blog_admin(request):
    """Admin panel for managing blog posts"""
    posts = BlogPost.objects.all().order_by('-created_at')
    
    # Filter by status
    status_filter = request.GET.get('status')
    if status_filter:
        posts = posts.filter(status=status_filter)
    
    # Search functionality
    query = request.GET.get('q')
    if query:
        posts = posts.filter(
            Q(title__icontains=query) |
            Q(author__username__icontains=query)
        )
    
    # Pagination
    paginator = Paginator(posts, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Compute plagiarism similarity summaries for posts on the current page
    # This runs automatically so admins can see results without running a check.
    similarity_summaries = {}
    try:
        # Compare each post on the page against all other posts with content
        other_posts_all = BlogPost.objects.exclude(content='')
        for current_post in page_obj:
            if not current_post.content:
                similarity_summaries[current_post.id] = []
                continue

            results = []
            for other_post in other_posts_all.exclude(pk=current_post.pk):
                # Use a fast containment-based similarity; do not expose algorithm details in UI
                try:
                    similarity_value = containment_similarity(current_post.content, other_post.content)
                except Exception:
                    similarity_value = 0.0
                if similarity_value > 5.0:
                    results.append({
                        'post': other_post,
                        'similarity_percentage': round(float(similarity_value), 2),
                    })

            # Sort by highest similarity and keep top 3 to keep UI concise
            results.sort(key=lambda r: r['similarity_percentage'], reverse=True)
            similarity_summaries[current_post.id] = results[:3]
    except Exception:
        # In case of any unexpected failure, do not break the admin page
        similarity_summaries = {}

    context = {
        'posts': page_obj,
        'status_filter': status_filter,
        'search_query': query,
        'similarity_summaries': similarity_summaries,
    }
    return render(request, 'pc/blog_admin.html', context)

@user_passes_test(is_admin)
def blog_approve(request, slug):
    """Approve a pending blog post"""
    post = get_object_or_404(BlogPost, slug=slug)
    post.status = 'published'
    post.save()
    messages.success(request, f'Blog post "{post.title}" has been approved and published!')
    return redirect('blog_admin')

@user_passes_test(is_admin)
def blog_reject(request, slug):
    """Reject a pending blog post"""
    post = get_object_or_404(BlogPost, slug=slug)
    post.status = 'rejected'
    post.save()
    messages.success(request, f'Blog post "{post.title}" has been rejected.')
    return redirect('blog_admin')

@user_passes_test(is_admin)
def blog_delete(request, slug):
    """Delete a blog post (admin only)"""
    post = get_object_or_404(BlogPost, slug=slug)
    if request.method == 'POST':
        post.delete()
        messages.success(request, f'Blog post "{post.title}" has been deleted.')
        return redirect('blog_admin')
    
    context = {
        'post': post,
    }
    return render(request, 'pc/blog_confirm_delete.html', context)

@user_passes_test(is_admin)
def blog_category_manage(request):
    """Manage blog categories"""
    if request.method == 'POST':
        form = BlogCategoryForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Category created successfully!')
            return redirect('blog_category_manage')
    else:
        form = BlogCategoryForm()
    
    categories = BlogCategory.objects.all()
    
    context = {
        'form': form,
        'categories': categories,
    }
    return render(request, 'pc/blog_category_manage.html', context)

@login_required
def blog_compare(request):
    """Compare blog posts using plagiarism detection algorithms"""
    
    # Pre-compute similarity matrix for all existing posts (like admin panel)
    existing_posts_similarities = []
    debug_info = {'total_posts': 0, 'posts_with_content': 0, 'comparisons_made': 0, 'similarities_found': 0}
    
    try:
        all_posts = BlogPost.objects.filter(status='published').exclude(content='')
        debug_info['total_posts'] = BlogPost.objects.filter(status='published').count()
        debug_info['posts_with_content'] = all_posts.count()
        
        print(f"DEBUG: Found {debug_info['posts_with_content']} published posts with content")
        
        # Test the algorithm with identical text
        if debug_info['posts_with_content'] >= 2:
            try:
                from plagiarismchecker.algorithm.fileSimilarity import findFileSimilarity
                test_posts = list(all_posts[:2])
                test_similarity = findFileSimilarity(test_posts[0].content, test_posts[0].content)
                print(f"DEBUG: Self-similarity test (should be ~100%): {test_similarity}%")
                
                cross_similarity = findFileSimilarity(test_posts[0].content, test_posts[1].content)
                print(f"DEBUG: Cross-similarity test: {cross_similarity}%")
            except Exception as te:
                print(f"DEBUG: Test error: {te}")
        
        # Create similarity pairs for all posts
        for i, current_post in enumerate(all_posts):
            if not current_post.content:
                continue
            
            print(f"DEBUG: Processing post '{current_post.title}' (content length: {len(current_post.content)})")
            print(f"DEBUG: Content preview: '{current_post.content[:100]}...'")
                
            post_similarities = []
            for other_post in all_posts.exclude(pk=current_post.pk):
                debug_info['comparisons_made'] += 1
                try:
                    # Use the same algorithm that works in manual comparison
                    from plagiarismchecker.algorithm.fileSimilarity import findFileSimilarity
                    similarity_value = findFileSimilarity(current_post.content, other_post.content)
                    print(f"DEBUG: Comparing '{current_post.title[:30]}...' vs '{other_post.title[:30]}...' = {similarity_value}%")
                    
                    # Also try containment similarity for comparison
                    try:
                        containment_value = containment_similarity(current_post.content, other_post.content)
                        print(f"DEBUG: Containment similarity = {containment_value}%")
                    except Exception as ce:
                        print(f"DEBUG: Containment error: {ce}")
                        
                except Exception as e:
                    print(f"DEBUG: Error computing similarity: {e}")
                    similarity_value = 0.0
                
                # Lower threshold for debugging - show similarities >1% instead of >5%
                if similarity_value > 1.0:  
                    debug_info['similarities_found'] += 1
                    post_similarities.append({
                        'post': other_post,
                        'similarity_percentage': round(float(similarity_value), 2),
                    })
            
            # Sort by highest similarity and keep top 3 for each post
            post_similarities.sort(key=lambda r: r['similarity_percentage'], reverse=True)
            
            if post_similarities:  # Only include posts that have similarities
                existing_posts_similarities.append({
                    'current_post': current_post,
                    'similarities': post_similarities[:3]  # Top 3 most similar
                })
                print(f"DEBUG: Post '{current_post.title}' has {len(post_similarities)} similar posts")
        
        print(f"DEBUG: Final results - {len(existing_posts_similarities)} posts with similarities")
        
    except Exception as e:
        print(f"DEBUG: Exception in similarity computation: {e}")
        existing_posts_similarities = []
    
    if request.method == 'POST':
        # Get the text to compare
        text_to_check = request.POST.get('text_to_check', '').strip()
        algorithm = request.POST.get('algorithm', 'cosine')
        
        if not text_to_check:
            messages.error(request, 'Please enter some text to check.')
            return redirect('blog_compare')
        
        # Get all published blog posts for comparison
        all_posts = BlogPost.objects.filter(status='published').exclude(content='')
        
        comparison_results = []
        
        for post in all_posts:
            post_content = post.content
            
            if algorithm == 'cosine':
                # Use cosine similarity (TF-IDF based)
                from plagiarismchecker.algorithm.fileSimilarity import findFileSimilarity
                similarity_percentage = findFileSimilarity(text_to_check, post_content)
            elif algorithm == 'containment':
                # Use containment similarity
                from plagiarismchecker.algorithm.main import containment_similarity
                similarity_percentage = containment_similarity(text_to_check, post_content)
            else:
                # Default to cosine similarity
                from plagiarismchecker.algorithm.fileSimilarity import findFileSimilarity
                similarity_percentage = findFileSimilarity(text_to_check, post_content)
            
            if similarity_percentage > 5:  # Only show results with >5% similarity
                # Find matching text segments
                matching_spans = find_plagiarized_spans(text_to_check, post_content, min_match_length=3)
                matching_segments = []
                
                for start, end in matching_spans[:3]:  # Limit to top 3 matches per source
                    user_text = text_to_check[start:end].strip()
                    # Find corresponding text in source post
                    source_spans = find_plagiarized_spans(post_content, text_to_check, min_match_length=3)
                    source_text = ""
                    if source_spans:
                        src_start, src_end = source_spans[0]
                        source_text = post_content[src_start:src_end].strip()
                    
                    if user_text and len(user_text) > 10:  # Only meaningful segments
                        matching_segments.append({
                            'user_text': user_text,
                            'source_text': source_text or user_text,
                            'char_start': start,
                            'char_end': end
                        })
                
                comparison_results.append({
                    'post': post,
                    'similarity_percentage': round(similarity_percentage, 2),
                    'algorithm_used': algorithm,
                    'matching_segments': matching_segments
                })
        
        # Sort by similarity percentage (highest first)
        comparison_results.sort(key=lambda x: x['similarity_percentage'], reverse=True)
        
        context = {
            'text_to_check': text_to_check,
            'algorithm': algorithm,
            'comparison_results': comparison_results,
            'total_posts_checked': len(all_posts),
            'posts_with_similarity': len(comparison_results),
            'existing_posts_similarities': existing_posts_similarities,
        }
        return render(request, 'pc/blog_compare.html', context)
    
    # GET request - show the form with pre-computed similarities
    context = {
        'total_posts': BlogPost.objects.filter(status='published').count(),
        'existing_posts_similarities': existing_posts_similarities,
        'debug_info': debug_info,  # Add debug info to context
    }
    return render(request, 'pc/blog_compare.html', context) 